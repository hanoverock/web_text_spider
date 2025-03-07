import requests
from bs4 import BeautifulSoup
from docx import Document

#输入网址
url=input('enter URL here  ')
headers={
    'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/95.0.4638.54 Safari/537.36'
    }
res=requests.get(url,headers=headers)

html=res.text
soup=BeautifulSoup(html,'html.parser')

#用class确定文本所在位置
anchor_class=input('enter the class value as text anchor  ')
block=soup.find(class_=anchor_class).find_all('p')

doc=Document()

#网址放在doc最上面
p=doc.add_paragraph(url+'\n'+'\n')

for item in block:
    p=doc.add_paragraph(item.text+'\n')

#命名文件
doc_name=input('name the doc as:  ')
doc.save(doc_name+'.docx')
    
